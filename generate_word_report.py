# -*- coding: utf-8 -*-
# generate_word_report.py
import sys
import os
import io

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import generate_charts as gc

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml as oxml
from lxml import etree

print("Generating charts for Word report...")
charts_bytes = {}
for name, fn in gc.CHARTS.items():
    charts_bytes[name] = gc.to_bytes(fn())
    print(f"  Chart '{name}' done.")

# ===== HELPERS =====

def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    h = hex_color.lstrip('#')
    shading_elm = parse_xml(
        '<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(
            nsdecls('w'), h))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = oxml.OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV', 'left', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = oxml.OxmlElement(tag)
            for k,v in kwargs[edge].items():
                element.set(qn('w:{}'.format(k)), v)
            tcBorders.append(element)

def add_style(doc, style_name, base_name, font_name, font_size, bold=False,
              color_hex=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base_name] if base_name in [s.name for s in styles] else None
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    if color_hex:
        font.color.rgb = hex_to_rgb(color_hex)
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return style

def add_para(doc, text, style_name, bold=False, italic=False, color_hex=None, alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = hex_to_rgb(color_hex)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_heading(doc, text, level, style_name):
    p = doc.add_paragraph(style=style_name)
    p.add_run(text)
    return p

def add_chart(doc, chart_name, width_inches=4.5, caption=None):
    buf = charts_bytes[chart_name]
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption, style='WASH_Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_styled(doc, headers, rows, caption=None, col_widths=None):
    if caption:
        cp = doc.add_paragraph(caption, style='WASH_Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.runs[0].bold = True
        cp.runs[0].font.color.rgb = hex_to_rgb('#1F3864')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '#0054A6')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = '#F2F7FF' if ri % 2 == 0 else '#FFFFFF'
        for ci, cell_val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

    if col_widths:
        for i, cw in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(cw)

    doc.add_paragraph()
    return table

def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_kpi_table(doc, kpis):
    """kpis = list of (label, value, sub)"""
    n = len(kpis)
    cols = min(3, n)
    rows_needed = (n + cols - 1) // cols
    table = doc.add_table(rows=rows_needed * 2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    colors = ['#0054A6','#C00000','#ED7D31','#FFC000','#00B0B9','#70AD47']
    idx = 0
    for r in range(rows_needed):
        for c in range(cols):
            if idx >= n: break
            label, value, sub = kpis[idx]
            color = colors[idx % len(colors)]
            cell_top = table.cell(r*2, c)
            cell_bot = table.cell(r*2+1, c)
            set_cell_bg(cell_top, color)
            set_cell_bg(cell_bot, '#F5F6FA')
            # value
            p = cell_top.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.bold = True; run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.name = 'Calibri Light'
            # label + sub
            p2 = cell_bot.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p2.add_run(label)
            r1.bold = True; r1.font.size = Pt(9)
            r1.font.name = 'Calibri'
            if sub:
                p2.add_run('\n' + sub).font.size = Pt(8)
            idx += 1

    doc.add_paragraph()

# ===== BUILD DOCUMENT =====

doc = Document()

# Page setup
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ===== DEFINE CUSTOM STYLES =====
print("Defining styles...")
add_style(doc, 'WASH_Title', 'Normal', 'Calibri Light', 36, bold=True,
          color_hex='#0054A6', alignment=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=12, space_after=12)
add_style(doc, 'WASH_H1', 'Normal', 'Calibri Light', 20, bold=True,
          color_hex='#0054A6', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=14, space_after=6)
add_style(doc, 'WASH_H2', 'Normal', 'Calibri', 14, bold=True,
          color_hex='#00B0B9', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=10, space_after=4)
add_style(doc, 'WASH_H3', 'Normal', 'Calibri', 12, bold=True,
          color_hex='#1F3864', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=8, space_after=4)
add_style(doc, 'WASH_Body', 'Normal', 'Calibri', 11, bold=False,
          color_hex='#222222', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=0, space_after=6)
add_style(doc, 'WASH_Caption', 'Normal', 'Calibri', 9, bold=False,
          color_hex='#595959', alignment=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=2, space_after=8)
add_style(doc, 'WASH_Bullet', 'Normal', 'Calibri', 10, bold=False,
          color_hex='#333333', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=3)

# ===== COVER PAGE =====
print("Building cover page...")

# Set first section header/footer off
section.different_first_page_header_footer = True
section.first_page_header.paragraphs[0].text = ''

# Cover background table
cover_table = doc.add_table(rows=1, cols=1)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cell = cover_table.cell(0, 0)
set_cell_bg(cover_cell, '#1F3864')
cover_cell.width = Cm(16)

p_spacer = cover_cell.paragraphs[0]
p_spacer.add_run('\n')

p_title = cover_cell.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('RAPPORT WASH')
r.bold = True; r.font.size = Pt(36); r.font.name = 'Calibri Light'
r.font.color.rgb = RGBColor(255,255,255)

p_sub = cover_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Eau, Assainissement & Hygiene - Cote d\'Ivoire')
r2.font.size = Pt(18); r2.font.name = 'Calibri Light'
r2.font.color.rgb = RGBColor(0,176,185)
r2.italic = True

p_sep = cover_cell.add_paragraph()
p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_sep.add_run('_' * 40)
r3.font.color.rgb = RGBColor(0,176,185)
r3.font.size = Pt(12)

p_meta = cover_cell.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_meta.add_run('Regions Goh & Loh Djiboua  |  Janvier-Février 2026  |  N = 1 000 ménages')
r4.font.size = Pt(11); r4.font.name = 'Calibri'
r4.font.color.rgb = RGBColor(200,215,230)

p_spacer2 = cover_cell.add_paragraph()
p_spacer2.add_run('\n\n')

p_stats = cover_cell.add_paragraph()
p_stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p_stats.add_run('1 000 ménages  |  166 variables  |  2 départements  |  8 sous-préfectures')
r5.font.size = Pt(10); r5.font.name = 'Calibri'
r5.font.color.rgb = RGBColor(180,200,220)

p_spacer3 = cover_cell.add_paragraph()
p_spacer3.add_run('\n')

doc.add_paragraph()

# Page break after cover
add_page_break(doc)

# ===== SETUP HEADER/FOOTER for subsequent pages =====
# Use the default (non-first) header
header = section.header
header_para = header.paragraphs[0]
header_para.clear()
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = header_para.add_run('Rapport WASH - Cote d\'Ivoire 2026')
run_h.font.name = 'Calibri'
run_h.font.size = Pt(9)
run_h.font.color.rgb = hex_to_rgb('#0054A6')
run_h.italic = True

# Footer with page numbers
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.clear()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add "Page X" using field
from docx.oxml.ns import qn
run_f1 = footer_para.add_run('Page ')
run_f1.font.size = Pt(9)
run_f1.font.name = 'Calibri'
run_f1.font.color.rgb = hex_to_rgb('#595959')

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run_fld = OxmlElement('w:r')
run_fld.append(fldChar1)
footer_para._p.append(run_fld)

instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
run_instr = OxmlElement('w:r')
run_instr.append(instrText)
footer_para._p.append(run_instr)

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run_fld2 = OxmlElement('w:r')
run_fld2.append(fldChar2)
footer_para._p.append(run_fld2)

run_f2 = footer_para.add_run(' | Enquête WASH 2026 | Regions Goh & Loh Djiboua')
run_f2.font.size = Pt(9)
run_f2.font.name = 'Calibri'
run_f2.font.color.rgb = hex_to_rgb('#595959')

# ===== SECTION 1: RESUME EXECUTIF =====
print("Section 1: Résumé exécutif...")
add_heading(doc, 'Section 1 - Résumé Exécutif', 0, 'WASH_H1')

add_para(doc,
    "Cette enquête WASH a ete conduite auprès de 1 000 ménages repartis dans les départements "
    "du Goh (570) et de Loh Djiboua (430) en Côte d'Ivoire, couvrant 8 sous-préfectures entre "
    "janvier et février 2026. Les donnees collectent 166 indicateurs sur l'eau, l'assainissement, "
    "la gestion des déchets, la pollution atmospherique et la disponibilité des ressources naturelles.",
    'WASH_Body')

add_heading(doc, '1.1 Indicateurs Clés (KPIs)', 0, 'WASH_H2')
add_kpi_table(doc, [
    ('Ménages enquêtes', '1 000', 'Jan (358) + Fev (642)'),
    ('Eau mauvaise qualité', '40%', '400 ménages'),
    ('Eau non traitee', '70.9%', '709 ménages'),
    ('Sans latrine', '33.6%', '336 ménages'),
    ('Pollution atmospherique', '49.4%', '494 signalements'),
    ('Eau suffisanté', '58%', '580 ménages satisfaits'),
])

add_heading(doc, '1.2 Alertes Principales', 0, 'WASH_H2')
alertes = [
    ('ALERTE CRITIQUE - Qualité de l\'eau',
     '40% des ménages ont une eau de mauvaise qualité et 70,9% ne traitent pas leur eau avant consommation. Risque élevé de maladies hydriques.'),
    ('ALERTE CRITIQUE - Assainissement',
     '33,6% des ménages n\'ont pas de latrines. La couverture en latrines améliorées reste très faible (9%).'),
    ('PREOCCUPATION - Maladies hydriques',
     '27,1% des ménages déclarent des maladies liées a l\'eau. La diarrhee domine (~180 cas) suivie de la typhoide (~120 cas).'),
    ('PREOCCUPATION - Ressources naturelles',
     '57,7% des ménages signalent une dégradation des ressources. 38,4% les jugent rares.'),
    ('POINT POSITIF - Disponibilité en eau',
     '58% des ménages disposent d\'une quantite d\'eau suffisanté pour leurs besoins quotidiens.'),
]
for title, text in alertes:
    p = doc.add_paragraph(style='WASH_Bullet')
    run = p.add_run('>> ' + title + ': ')
    run.bold = True; run.font.size = Pt(10)
    run.font.color.rgb = hex_to_rgb('#0054A6')
    p.add_run(text).font.size = Pt(10)

doc.add_paragraph()

# ===== SECTION 2: CONTEXTE & METHODOLOGIE =====
print("Section 2: Contexte...")
add_page_break(doc)
add_heading(doc, 'Section 2 - Contexte & Méthodologie', 0, 'WASH_H1')

add_para(doc,
    "L'enquête WASH 2026 vise a evaluer les conditions d'accès a l'eau potable, l'état de "
    "l'assainissement et les pratiques d'hygiene dans les regions du Goh et de Loh Djiboua. "
    "Ces deux départements presentent des enjeux importants en matiere de santé publique et "
    "de gestion environnementale.",
    'WASH_Body')

add_heading(doc, '2.1 Répartition par département', 0, 'WASH_H2')
add_chart(doc, 'departements', width_inches=3.5,
          caption='Figure 1 - Répartition de l\'echantillon par département (N=1000)')

add_heading(doc, '2.2 Répartition par sous-préfecture', 0, 'WASH_H2')
add_table_styled(doc,
    headers=['Sous-préfecture', 'Departement', 'Enquêtes', '% echantillon'],
    rows=[
        ['Gagnoa', 'Goh', '~115', '~11.5%'],
        ['Bayota', 'Goh', '~95', '~9.5%'],
        ['Ouragahio', 'Goh', '~90', '~9.0%'],
        ['Guitry', 'Goh', '~88', '~8.8%'],
        ['Guiberoua', 'Goh', '~85', '~8.5%'],
        ['Divo', 'Loh Djiboua', '~155', '~15.5%'],
        ['Lakota', 'Loh Djiboua', '~140', '~14.0%'],
        ['Hire', 'Loh Djiboua', '~135', '~13.5%'],
        ['TOTAL', '2 départements', '1 000', '100%'],
    ],
    caption='Tableau 1 - Répartition par sous-préfecture',
    col_widths=[5.0, 4.0, 3.5, 3.5]
)

add_table_styled(doc,
    headers=['Thematique', 'Variables clés', 'Type'],
    rows=[
        ['Eau potable', 'Source, qualité, temps accès, traitement', 'Quantitatif'],
        ['Santé hydrique', 'Maladies liées a l\'eau, frequence, type', 'Mixte'],
        ['Assainissement', 'Type de latrine, pratiques hygiene', 'Quantitatif'],
        ['Déchets', 'Mode de gestion, satisfaction, exposition', 'Quantitatif'],
        ['Environnement', 'Pollution, qualité air, nuisances sonores', 'Qualitatif'],
        ['Ressources', 'Disponibilité, dégradation, respect', 'Qualitatif'],
    ],
    caption='Tableau 2 - Outils et variables de collecte',
    col_widths=[4.0, 7.0, 3.0]
)

# ===== SECTION 3: EAU POTABLE =====
print("Section 3: Eau potable...")
add_page_break(doc)
add_heading(doc, 'Section 3 - Eau Potable', 0, 'WASH_H1')

add_para(doc,
    "L'accès a l'eau potable est l'un des defis majeurs identifies dans cette enquête. "
    "Si 58% des ménages disposent d'une eau en quantite suffisanté, la qualité reste "
    "preoccupante avec 40% d'eau de mauvaise qualité et seulement 29,1% des ménages "
    "qui traitent leur eau.",
    'WASH_Body')

add_heading(doc, '3.1 Sources d\'approvisionnement', 0, 'WASH_H2')
add_chart(doc, 'sources', width_inches=4.5,
          caption='Figure 2 - Sources d\'approvisionnement en eau (N=1000)')

add_table_styled(doc,
    headers=['Source', 'Nb ménages', '%', 'Risque'],
    rows=[
        ['Puits traditionnel', '275', '27.5%', 'Eleve'],
        ['Puits amélioré', '200', '20.0%', 'Modere'],
        ['Forage', '196', '19.6%', 'Faible'],
        ['Riviere / Source', '167', '16.7%', 'Très élevé'],
        ['Autrès', '82', '8.2%', 'Indetermine'],
        ['SODECI (reseau)', '80', '8.0%', 'Faible'],
        ['Total', '1 000', '100%', '-'],
    ],
    caption='Tableau 3 - Detail des sources d\'eau',
    col_widths=[5.5, 3.5, 3.0, 4.0]
)

add_heading(doc, '3.2 Qualité de l\'eau par département', 0, 'WASH_H2')
add_chart(doc, 'qualite_dept', width_inches=4.5,
          caption='Figure 3 - Qualité de l\'eau percue par département')

add_chart(doc, 'traitement', width_inches=3.5,
          caption='Figure 4 - Traitement de l\'eau au niveau domestique')

add_table_styled(doc,
    headers=['Departement', 'Bonne', 'Acceptable', 'Mauvaise', 'Total', '% mauvaise'],
    rows=[
        ['Goh', '118 (20.7%)', '221 (38.8%)', '231 (40.5%)', '570', '40.5%'],
        ['Loh Djiboua', '90 (20.9%)', '171 (39.8%)', '169 (39.3%)', '430', '39.3%'],
        ['Total', '208 (20.8%)', '392 (39.2%)', '400 (40.0%)', '1 000', '40.0%'],
    ],
    caption='Tableau 4 - Qualité de l\'eau par département',
    col_widths=[3.5, 3.0, 3.0, 3.0, 2.0, 2.5]
)

add_heading(doc, '3.3 Accès et maladies hydriques', 0, 'WASH_H2')
add_chart(doc, 'acces', width_inches=4.5,
          caption='Figure 5 - Temps d\'accès a la source d\'eau')

add_chart(doc, 'maladies', width_inches=4.5,
          caption='Figure 6 - Maladies hydriques déclarees par département')

add_table_styled(doc,
    headers=['Tranche de temps', 'Ménages', '%', 'Norme OMS', 'Statut'],
    rows=[
        ['Moins de 30 minutes', '386', '38.6%', 'Conforme', 'Acceptable'],
        ['30 a 60 minutes', '341', '34.1%', 'Limite', 'Attention'],
        ['Plus de 60 minutes', '273', '27.3%', 'Non conforme', 'Critique'],
    ],
    caption='Tableau 5 - Temps d\'accès a l\'eau et implications',
    col_widths=[4.5, 2.5, 2.5, 3.5, 3.0]
)

add_table_styled(doc,
    headers=['Maladie', 'Goh', 'Loh Djiboua', 'Total estime', 'Gravite'],
    rows=[
        ['Diarrhee', '~115', '~87', '~202', 'Moderee'],
        ['Typhoide', '~72', '~55', '~127', 'Haute'],
        ['Hepatite A', '~60', '~43', '~103', 'Haute'],
        ['Cholera', '~50', '~34', '~84', 'Très haute'],
        ['Total déclarés', '~297', '~219', '~516 cas', '-'],
    ],
    caption='Tableau 6 - Maladies hydriques déclarees',
    col_widths=[4.0, 2.5, 3.0, 3.0, 3.5]
)

# ===== SECTION 4: ASSAINISSEMENT =====
print("Section 4: Assainissement...")
add_page_break(doc)
add_heading(doc, 'Section 4 - Assainissement', 0, 'WASH_H1')

add_para(doc,
    "L'assainissement represente un defi majeur dans les deux départements. Pres d'un tiers "
    "des ménages (33,6%) n'ont pas accès a des latrines, et seulement 14,8% disposent de latrines "
    "améliorées ou de WC modernes, loin des objectifs de l'ODD 6.",
    'WASH_Body')

add_chart(doc, 'latrines', width_inches=5.0,
          caption='Figure 7 - Types d\'infrastructures sanitaires (N=1000)')

add_table_styled(doc,
    headers=['Type d\'infrastructure', 'Nb ménages', '%', 'Conformite ODD 6'],
    rows=[
        ['Latrine simple', '424', '42.4%', 'Partielle'],
        ['Aucune latrine (defecation a l\'air libre)', '336', '33.6%', 'Non conforme'],
        ['Autre type', '92', '9.2%', 'Indetermine'],
        ['Latrines améliorées', '90', '9.0%', 'Conforme'],
        ['WC modernes', '58', '5.8%', 'Conforme'],
        ['Total', '1 000', '100%', '-'],
    ],
    caption='Tableau 7 - Types de latrines et conformite ODD 6',
    col_widths=[6.0, 3.0, 2.5, 4.5]
)

add_table_styled(doc,
    headers=['Indicateur', 'Valeur', 'Cible ODD 6', 'Ecart'],
    rows=[
        ['Accès a des latrines (tout type)', '66.4%', '100%', '-33.6 pts'],
        ['Latrines améliorées + WC', '14.8%', '100%', '-85.2 pts'],
        ['Defecation a l\'air libre (DAL)', '33.6%', '0%', '+33.6 pts'],
        ['Traitement de l\'eau domestique', '29.1%', '100%', '-70.9 pts'],
    ],
    caption='Tableau 8 - Indicateurs d\'assainissement vs cibles ODD 6',
    col_widths=[5.5, 3.0, 3.0, 4.5]
)

# ===== SECTION 5: DECHETS =====
print("Section 5: Déchets...")
add_page_break(doc)
add_heading(doc, 'Section 5 - Gestion des Déchets', 0, 'WASH_H1')

add_para(doc,
    "La gestion des déchets est une problematique complexe dans les deux départements. "
    "52,3% des ménages sont exposes a des decharges sauvages. La satisfaction reste majoritairement "
    "negative avec 62% de ménages peu ou pas satisfaits de la gestion actuelle.",
    'WASH_Body')

add_chart(doc, 'gestion_modes', width_inches=5.0,
          caption='Figure 8 - Modes de gestion des déchets')

add_chart(doc, 'dechets_sat', width_inches=5.0,
          caption='Figure 9 - Satisfaction quant a la gestion des déchets')

add_table_styled(doc,
    headers=['Mode de gestion', 'Ménages', '%', 'Impact environnemental'],
    rows=[
        ['Collecte / Valorisation', '515', '51.5%', 'Positif'],
        ['Abandon / Decharge sauvage', '501', '50.1%', 'Très negatif'],
        ['Brulage', '499', '49.9%', 'Negatif'],
        ['Reutilisation', '287', '28.7%', 'Positif'],
    ],
    caption='Tableau 9 - Modes de gestion des déchets',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Niveau de satisfaction', 'Ménages', '%', 'Tendance'],
    rows=[
        ['Peu satisfait', '378', '37.8%', 'Negatif'],
        ['Satisfait', '271', '27.1%', 'Positif'],
        ['Pas satisfait', '242', '24.2%', 'Negatif'],
        ['Très satisfait', '109', '10.9%', 'Positif'],
        ['Total', '1 000', '100%', '-'],
    ],
    caption='Tableau 10 - Satisfaction gestion des déchets',
    col_widths=[5.0, 3.0, 2.5, 5.5]
)

# ===== SECTION 6: POLLUTION =====
print("Section 6: Pollution...")
add_page_break(doc)
add_heading(doc, 'Section 6 - Pollution & Qualité de l\'Air', 0, 'WASH_H1')

add_para(doc,
    "La pollution atmospherique est signalee par pres d'un ménage sur deux (49,4%). "
    "La qualité de l'air est jugee mauvaise ou très mauvaise par 51,6% des repondants. "
    "Les nuisances sonores affectent egalement une majorite des ménages (57,3%).",
    'WASH_Body')

add_chart(doc, 'pollution', width_inches=4.5,
          caption='Figure 10 - Pollution atmospherique signalee par département')

add_chart(doc, 'qualite_air', width_inches=4.5,
          caption='Figure 11 - Qualité de l\'air percue (N=1000)')

add_table_styled(doc,
    headers=['Indicateur', 'Ménages', '%', 'Niveau de risque'],
    rows=[
        ['Pollution atmospherique signalee', '494', '49.4%', 'Preoccupant'],
        ['Qualité air : Mauvaise', '364', '36.4%', 'Critique'],
        ['Qualité air : Moyenne', '361', '36.1%', 'Modere'],
        ['Qualité air : Très mauvaise', '152', '15.2%', 'Très critique'],
        ['Qualité air : Bonne', '123', '12.3%', 'Acceptable'],
        ['Nuisances sonores', '573', '57.3%', 'Preoccupant'],
    ],
    caption='Tableau 11 - Indicateurs de pollution et qualité de l\'air',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Departement', 'Pollution signalee', 'Pas de pollution', 'Taux'],
    rows=[
        ['Goh', '285', '285', '50.0%'],
        ['Loh Djiboua', '209', '221', '48.6%'],
        ['Total', '494', '506', '49.4%'],
    ],
    caption='Tableau 12 - Pollution par département',
    col_widths=[4.0, 4.0, 4.0, 4.0]
)

# ===== SECTION 7: RESSOURCES NATURELLES =====
print("Section 7: Ressources naturelles...")
add_page_break(doc)
add_heading(doc, 'Section 7 - Ressources Naturelles', 0, 'WASH_H1')

add_para(doc,
    "La situation des ressources naturelles est alarmante : 57,7% des ménages signalent "
    "une dégradation des ressources, et seulement 12,1% les considerent abondantes. "
    "Cette situation indique une prise de conscience croissanté mais des pratiques qui "
    "restent insuffisantés pour enrayer la tendance.",
    'WASH_Body')

add_chart(doc, 'ressources', width_inches=5.0,
          caption='Figure 12 - Disponibilité percue des ressources naturelles (N=1000)')

add_table_styled(doc,
    headers=['Niveau de disponibilité', 'Ménages', '%', 'Statut'],
    rows=[
        ['Rares', '384', '38.4%', 'Critique'],
        ['Moyennement disponibles', '335', '33.5%', 'Precaire'],
        ['En voie de disparition', '160', '16.0%', 'Urgence'],
        ['Abondantes', '121', '12.1%', 'Favorable'],
    ],
    caption='Tableau 13 - Disponibilité des ressources naturelles',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Indicateur ressources', 'Valeur', 'Interprétation'],
    rows=[
        ['Dégradation signalee', '57.7% (577)', 'Situation critique'],
        ['Respect environnement : Fort', '35.8% (358)', 'Conscience presente'],
        ['Respect environnement : Moyen', '30.8% (308)', 'A renforcer'],
        ['Respect environnement : Très fort', '18.3% (183)', 'Minorite engagee'],
        ['Respect environnement : Faible', '15.1% (151)', 'Risque élevé'],
    ],
    caption='Tableau 14 - Indicateurs ressources naturelles et respect environnement',
    col_widths=[5.5, 4.0, 6.5]
)

# ===== SECTION 8: ANALYSE COMPARATIVE =====
print("Section 8: Analyse comparative...")
add_page_break(doc)
add_heading(doc, 'Section 8 - Analyse Comparative', 0, 'WASH_H1')

add_para(doc,
    "La comparaison des indicateurs WASH entre les deux départements revele des profils "
    "relativement similaires. Le département du Goh presente legerement plus de cas de "
    "maladies et une plus forte proportion d'eau traitee, tandis que Loh Djiboua montre "
    "un meilleur accès en moins de 30 minutes.",
    'WASH_Body')

add_chart(doc, 'comparaison', width_inches=5.5,
          caption='Figure 13 - Comparaison des indicateurs WASH clés par département')

add_table_styled(doc,
    headers=['Indicateur', 'Ensemble', 'Goh (N=570)', 'Loh Djiboua (N=430)', 'Ecart'],
    rows=[
        ['Eau bonne qualité', '20.8%', '20.7%', '20.9%', '+0.2 pts'],
        ['Eau mauvaise qualité', '40.0%', '40.5%', '39.3%', '-1.2 pts'],
        ['Accès <30 min', '38.6%', '37.5%', '40.0%', '+2.5 pts'],
        ['Accès >60 min', '27.3%', '28.1%', '26.3%', '-1.8 pts'],
        ['Eau traitee', '29.1%', '30.0%', '28.1%', '-1.9 pts'],
        ['Maladies hydriques', '27.1%', '26.7%', '27.7%', '+1.0 pts'],
        ['Sans latrine (DAL)', '33.6%', '33.3%', '33.3%', '0.0 pts'],
        ['Pollution atmospherique', '49.4%', '50.0%', '48.6%', '-1.4 pts'],
        ['Ressources rares', '38.4%', '~39.0%', '~37.5%', '-1.5 pts'],
        ['Dégradation ressources', '57.7%', '~58.0%', '~57.2%', '-0.8 pts'],
        ['Eau suffisanté', '58.0%', '~58.5%', '~57.2%', '-1.3 pts'],
        ['Exposes aux déchets', '52.3%', '~53.0%', '~51.4%', '-1.6 pts'],
    ],
    caption='Tableau 15 - Tableau comparatif complet des indicateurs WASH',
    col_widths=[4.0, 2.5, 3.0, 3.5, 3.0]
)

# ===== SECTION 9: RECOMMANDATIONS =====
print("Section 9: Recommandations...")
add_page_break(doc)
add_heading(doc, 'Section 9 - Recommandations', 0, 'WASH_H1')

add_para(doc,
    "Sur la base des resultats de cette enquête, les recommandations suivantes sont formulees, "
    "classees par ordre de priorite selon l'urgence et l'impact potentiel sur la santé publique "
    "et l'environnement.",
    'WASH_Body')

recos = [
    ('Priorite 1 - Qualité et traitement de l\'eau', [
        'Campagnes massives sur le traitement domestique de l\'eau (chloration, filtration, ebullition)',
        'Construction de forages supplementaires dans les zones a risque',
        'Tests reguliers de la qualité des sources d\'eau',
        'Rehabilitation des puits traditionnels existants',
    ]),
    ('Priorite 2 - Assainissement & Latrines', [
        'Programme d\'elimination de la defecation a l\'air libre (FDAL)',
        'Subventions pour la construction de latrines améliorées',
        'Sensibilisation aux maladies hydriques liées au manque d\'assainissement',
        'Formation des macons locaux a la construction de latrines',
    ]),
    ('Priorite 3 - Gestion des déchets', [
        'Mise en place de systemes de collecte structures dans les zones deficitaires',
        'Promotion du compostage et de la valorisation des déchets organiques',
        'Interdiction et sanctions du brulage a ciel ouvert',
        'Depots de transit reglementes et securises',
    ]),
    ('Priorite 4 - Pollution atmospherique', [
        'Cartographie et surveillance des sources de pollution',
        'Reglementation des emissions industrielles et artisanales',
        'Promotion de foyers améliorés pour reduire la fumee domestique',
        'Zones tampon autour des agglomerations',
    ]),
    ('Priorite 5 - Ressources naturelles', [
        'Plans locaux de gestion durable des ressources naturelles',
        'Reboisement et restauration des terres degradees',
        'Sensibilisation communautaire sur la preservation de l\'environnement',
        'Systemes de suivi participatif des ressources',
    ]),
    ('Priorite 6 - Suivi & Gouvernance', [
        'Systemes de suivi-évaluation des indicateurs WASH',
        'Renforcement des capacites des autorites locales',
        'Implication des communautés dans la gestion des infrastructures',
        'Repetition de l\'enquête tous les 2 ans pour mesurer les progres',
    ]),
]

for title, items in recos:
    add_heading(doc, title, 0, 'WASH_H2')
    for item in items:
        p = doc.add_paragraph(style='WASH_Bullet')
        run = p.add_run('   >> ' + item)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    doc.add_paragraph()

# ===== SECTION 10: CONCLUSION =====
print("Section 10: Conclusion...")
add_page_break(doc)
add_heading(doc, 'Section 10 - Conclusion', 0, 'WASH_H1')

add_para(doc,
    "Cette enquête WASH conduite auprès de 1 000 ménages dans les départements du Goh et "
    "de Loh Djiboua en Côte d'Ivoire met en lumiere des defis considerables en matiere "
    "d'eau potable, d'assainissement et de gestion environnementale.",
    'WASH_Body')

add_para(doc,
    "Les resultats les plus preoccupants concernent la qualité de l'eau (40% de mauvaise "
    "qualité), le faible taux de traitement domestique (29,1%), et la prevalence de la "
    "defecation a l'air libre (33,6%). Ces facteurs combines expliquent en grande partie "
    "les 27,1% de ménages declarant des maladies liées a l'eau.",
    'WASH_Body')

add_para(doc,
    "Par rapport aux Objectifs de Developpement Durable (ODD 6) - garantir l'accès a l'eau "
    "potable et a l'assainissement a tous - les deux départements se trouvent encore loin des "
    "cibles. Des investissements cibles et des programmes de changement de comportement sont "
    "indispensables pour reduire les inegalites d'accès et proteger la santé des populations.",
    'WASH_Body')

add_para(doc,
    "Les profils des deux départements etant très similaires, les interventions peuvent etre "
    "planifiees de maniere coordonnee, tout en tenant compte des specificites locales de chaque "
    "sous-préfecture. Une repetition de l'enquête dans deux ans permettra de mesurer les progres "
    "accomplis et d'adapter les stratégies d'intervention.",
    'WASH_Body')

add_heading(doc, 'Synthese des indicateurs clés', 0, 'WASH_H2')
add_table_styled(doc,
    headers=['Dimension', 'Indicateur principal', 'Valeur', 'Priorite'],
    rows=[
        ['Eau', 'Qualité mauvaise', '40%', 'Critique'],
        ['Eau', 'Non traite', '70.9%', 'Critique'],
        ['Accès', 'Plus de 60 min', '27.3%', 'Eleve'],
        ['Santé', 'Maladies hydriques', '27.1%', 'Eleve'],
        ['Assainissement', 'Sans latrine', '33.6%', 'Critique'],
        ['Déchets', 'Peu/pas satisfait', '62.0%', 'Modere'],
        ['Pollution', 'Signalee', '49.4%', 'Modere'],
        ['Ressources', 'Dégradation', '57.7%', 'Eleve'],
        ['Environnement', 'Ressources rares', '38.4%', 'Eleve'],
    ],
    caption='Tableau 16 - Synthese des indicateurs clés et niveaux de priorite',
    col_widths=[4.0, 5.5, 3.0, 3.5]
)

# ===== SAVE =====
out_path = r'C:\Users\ITrebi\Music\PortFolio Upwork\Excel Dashboard\WASH_Rapport.docx'
doc.save(out_path)
print("Word report saved to: " + out_path)
