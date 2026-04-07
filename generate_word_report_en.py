# -*- coding: utf-8 -*-
# generate_word_report_en.py
import sys
import os
import io

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import generate_charts_en as gc

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml as oxml
from lxml import etree

print("Generating charts for Word report...")
charts_bytes = {}
for name, fn in gc.CHARTS.items():
    charts_bytes[name] = gc.to_bytes(fn())
    print(f"  Chart '{name}' done.")

# ===== HELPERS =====

def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    h = hex_color.lstrip('#')
    shading_elm = parse_xml(
        '<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(
            nsdecls('w'), h))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = oxml.OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV', 'left', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = oxml.OxmlElement(tag)
            for k,v in kwargs[edge].items():
                element.set(qn('w:{}'.format(k)), v)
            tcBorders.append(element)

def add_style(doc, style_name, base_name, font_name, font_size, bold=False,
              color_hex=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base_name] if base_name in [s.name for s in styles] else None
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    if color_hex:
        font.color.rgb = hex_to_rgb(color_hex)
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return style

def add_para(doc, text, style_name, bold=False, italic=False, color_hex=None, alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = hex_to_rgb(color_hex)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_heading(doc, text, level, style_name):
    p = doc.add_paragraph(style=style_name)
    p.add_run(text)
    return p

def add_chart(doc, chart_name, width_inches=4.5, caption=None):
    buf = charts_bytes[chart_name]
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption, style='WASH_Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_styled(doc, headers, rows, caption=None, col_widths=None):
    if caption:
        cp = doc.add_paragraph(caption, style='WASH_Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.runs[0].bold = True
        cp.runs[0].font.color.rgb = hex_to_rgb('#1F3864')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '#0054A6')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = '#F2F7FF' if ri % 2 == 0 else '#FFFFFF'
        for ci, cell_val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

    if col_widths:
        for i, cw in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(cw)

    doc.add_paragraph()
    return table

def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_kpi_table(doc, kpis):
    """kpis = list of (label, value, sub)"""
    n = len(kpis)
    cols = min(3, n)
    rows_needed = (n + cols - 1) // cols
    table = doc.add_table(rows=rows_needed * 2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    colors = ['#0054A6','#C00000','#ED7D31','#FFC000','#00B0B9','#70AD47']
    idx = 0
    for r in range(rows_needed):
        for c in range(cols):
            if idx >= n: break
            label, value, sub = kpis[idx]
            color = colors[idx % len(colors)]
            cell_top = table.cell(r*2, c)
            cell_bot = table.cell(r*2+1, c)
            set_cell_bg(cell_top, color)
            set_cell_bg(cell_bot, '#F5F6FA')
            # value
            p = cell_top.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.bold = True; run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.name = 'Calibri Light'
            # label + sub
            p2 = cell_bot.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p2.add_run(label)
            r1.bold = True; r1.font.size = Pt(9)
            r1.font.name = 'Calibri'
            if sub:
                p2.add_run('\n' + sub).font.size = Pt(8)
            idx += 1

    doc.add_paragraph()

# ===== BUILD DOCUMENT =====

doc = Document()

# Page setup
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ===== DEFINE CUSTOM STYLES =====
print("Defining styles...")
add_style(doc, 'WASH_Title', 'Normal', 'Calibri Light', 36, bold=True,
          color_hex='#0054A6', alignment=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=12, space_after=12)
add_style(doc, 'WASH_H1', 'Normal', 'Calibri Light', 20, bold=True,
          color_hex='#0054A6', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=14, space_after=6)
add_style(doc, 'WASH_H2', 'Normal', 'Calibri', 14, bold=True,
          color_hex='#00B0B9', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=10, space_after=4)
add_style(doc, 'WASH_H3', 'Normal', 'Calibri', 12, bold=True,
          color_hex='#1F3864', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=8, space_after=4)
add_style(doc, 'WASH_Body', 'Normal', 'Calibri', 11, bold=False,
          color_hex='#222222', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=0, space_after=6)
add_style(doc, 'WASH_Caption', 'Normal', 'Calibri', 9, bold=False,
          color_hex='#595959', alignment=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=2, space_after=8)
add_style(doc, 'WASH_Bullet', 'Normal', 'Calibri', 10, bold=False,
          color_hex='#333333', alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=3)

# ===== COVER PAGE =====
print("Building cover page...")

# Set first section header/footer off
section.different_first_page_header_footer = True
section.first_page_header.paragraphs[0].text = ''

# Cover background table
cover_table = doc.add_table(rows=1, cols=1)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cell = cover_table.cell(0, 0)
set_cell_bg(cover_cell, '#1F3864')
cover_cell.width = Cm(16)

p_spacer = cover_cell.paragraphs[0]
p_spacer.add_run('\n')

p_title = cover_cell.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('WASH SURVEY REPORT')
r.bold = True; r.font.size = Pt(36); r.font.name = 'Calibri Light'
r.font.color.rgb = RGBColor(255,255,255)

p_sub = cover_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Water, Sanitation & Hygiene - Cote d\'Ivoire')
r2.font.size = Pt(18); r2.font.name = 'Calibri Light'
r2.font.color.rgb = RGBColor(0,176,185)
r2.italic = True

p_sep = cover_cell.add_paragraph()
p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_sep.add_run('_' * 40)
r3.font.color.rgb = RGBColor(0,176,185)
r3.font.size = Pt(12)

p_meta = cover_cell.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_meta.add_run('Goh & Loh Djiboua Regions  |  January-February 2026  |  N = 1,000 households')
r4.font.size = Pt(11); r4.font.name = 'Calibri'
r4.font.color.rgb = RGBColor(200,215,230)

p_spacer2 = cover_cell.add_paragraph()
p_spacer2.add_run('\n\n')

p_stats = cover_cell.add_paragraph()
p_stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p_stats.add_run('1,000 households  |  166 variables  |  2 departments  |  8 sub-prefectures')
r5.font.size = Pt(10); r5.font.name = 'Calibri'
r5.font.color.rgb = RGBColor(180,200,220)

p_spacer3 = cover_cell.add_paragraph()
p_spacer3.add_run('\n')

doc.add_paragraph()

# Page break after cover
add_page_break(doc)

# ===== SETUP HEADER/FOOTER for subsequent pages =====
header = section.header
header_para = header.paragraphs[0]
header_para.clear()
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = header_para.add_run('WASH Survey Report - Cote d\'Ivoire 2026')
run_h.font.name = 'Calibri'
run_h.font.size = Pt(9)
run_h.font.color.rgb = hex_to_rgb('#0054A6')
run_h.italic = True

# Footer with page numbers
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.clear()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

from docx.oxml.ns import qn
run_f1 = footer_para.add_run('Page ')
run_f1.font.size = Pt(9)
run_f1.font.name = 'Calibri'
run_f1.font.color.rgb = hex_to_rgb('#595959')

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run_fld = OxmlElement('w:r')
run_fld.append(fldChar1)
footer_para._p.append(run_fld)

instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
run_instr = OxmlElement('w:r')
run_instr.append(instrText)
footer_para._p.append(run_instr)

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run_fld2 = OxmlElement('w:r')
run_fld2.append(fldChar2)
footer_para._p.append(run_fld2)

run_f2 = footer_para.add_run(' | WASH Survey 2026 | Goh & Loh Djiboua Regions')
run_f2.font.size = Pt(9)
run_f2.font.name = 'Calibri'
run_f2.font.color.rgb = hex_to_rgb('#595959')

# ===== SECTION 1: EXECUTIVE SUMMARY =====
print("Section 1: Executive Summary...")
add_heading(doc, 'Section 1 - Executive Summary', 0, 'WASH_H1')

add_para(doc,
    "This WASH survey was conducted among 1,000 households distributed across the departments "
    "of Goh (570) and Loh Djiboua (430) in Cote d'Ivoire, covering 8 sub-prefectures between "
    "January and February 2026. The data collected covers 166 indicators on water, sanitation, "
    "waste management, atmospheric pollution and natural resources availability.",
    'WASH_Body')

add_heading(doc, '1.1 Key Performance Indicators (KPIs)', 0, 'WASH_H2')
add_kpi_table(doc, [
    ('Households surveyed', '1,000', 'Jan (358) + Feb (642)'),
    ('Poor water quality', '40%', '400 households'),
    ('Untreated water', '70.9%', '709 households'),
    ('No latrine', '33.6%', '336 households'),
    ('Atmospheric pollution', '49.4%', '494 reports'),
    ('Sufficient water', '58%', '580 satisfied households'),
])

add_heading(doc, '1.2 Key Alerts', 0, 'WASH_H2')
alerts = [
    ('CRITICAL ALERT - Water Quality',
     '40% of households have poor quality water and 70.9% do not treat their water before consumption. High risk of waterborne diseases.'),
    ('CRITICAL ALERT - Sanitation',
     '33.6% of households have no latrines. Coverage of improved latrines remains very low (9%).'),
    ('CONCERN - Waterborne Diseases',
     '27.1% of households report water-related diseases. Diarrhea dominates (~180 cases) followed by typhoid (~120 cases).'),
    ('CONCERN - Natural Resources',
     '57.7% of households report resource degradation. 38.4% consider them scarce, and 16% consider them near depletion.'),
    ('POSITIVE POINT - Water Availability',
     '58% of households have a sufficient quantity of water for their daily needs.'),
]
for title, text in alerts:
    p = doc.add_paragraph(style='WASH_Bullet')
    run = p.add_run('>> ' + title + ': ')
    run.bold = True; run.font.size = Pt(10)
    run.font.color.rgb = hex_to_rgb('#0054A6')
    p.add_run(text).font.size = Pt(10)

doc.add_paragraph()

# ===== SECTION 2: CONTEXT & METHODOLOGY =====
print("Section 2: Context & Methodology...")
add_page_break(doc)
add_heading(doc, 'Section 2 - Context & Methodology', 0, 'WASH_H1')

add_para(doc,
    "The 2026 WASH survey aims to assess the conditions of access to drinking water, the state of "
    "sanitation and hygiene practices in the Goh and Loh Djiboua regions. "
    "These two departments face significant challenges in terms of public health and "
    "environmental management.",
    'WASH_Body')

add_heading(doc, '2.1 Sample distribution by department', 0, 'WASH_H2')
add_chart(doc, 'departements', width_inches=3.5,
          caption='Figure 1 - Sample distribution by department (N=1000)')

add_heading(doc, '2.2 Distribution by sub-prefecture', 0, 'WASH_H2')
add_table_styled(doc,
    headers=['Sub-prefecture', 'Department', 'Surveys', '% of sample'],
    rows=[
        ['Gagnoa', 'Goh', '~115', '~11.5%'],
        ['Bayota', 'Goh', '~95', '~9.5%'],
        ['Ouragahio', 'Goh', '~90', '~9.0%'],
        ['Guitry', 'Goh', '~88', '~8.8%'],
        ['Guiberoua', 'Goh', '~85', '~8.5%'],
        ['Divo', 'Loh Djiboua', '~155', '~15.5%'],
        ['Lakota', 'Loh Djiboua', '~140', '~14.0%'],
        ['Hire', 'Loh Djiboua', '~135', '~13.5%'],
        ['TOTAL', '2 departments', '1,000', '100%'],
    ],
    caption='Table 1 - Distribution by sub-prefecture',
    col_widths=[5.0, 4.0, 3.5, 3.5]
)

add_table_styled(doc,
    headers=['Theme', 'Key variables', 'Type'],
    rows=[
        ['Drinking water', 'Source, quality, access time, treatment', 'Quantitative'],
        ['Water health', 'Water-related diseases, frequency, type', 'Mixed'],
        ['Sanitation', 'Latrine type, hygiene practices', 'Quantitative'],
        ['Waste', 'Disposal method, satisfaction, exposure', 'Quantitative'],
        ['Environment', 'Pollution, air quality, noise nuisances', 'Qualitative'],
        ['Resources', 'Availability, degradation, stewardship', 'Qualitative'],
    ],
    caption='Table 2 - Data collection tools and variables',
    col_widths=[4.0, 7.0, 3.0]
)

# ===== SECTION 3: WATER & SANITATION =====
print("Section 3: Water & Sanitation...")
add_page_break(doc)
add_heading(doc, 'Section 3 - Water & Sanitation', 0, 'WASH_H1')

add_para(doc,
    "Access to drinking water is one of the major challenges identified in this survey. "
    "While 58% of households have a sufficient quantity of water, quality remains "
    "concerning with 40% poor quality water and only 29.1% of households "
    "treating their water.",
    'WASH_Body')

add_heading(doc, '3.1 Water supply sources', 0, 'WASH_H2')
add_chart(doc, 'sources', width_inches=4.5,
          caption='Figure 2 - Water supply sources (N=1000)')

add_table_styled(doc,
    headers=['Source', 'No. households', '%', 'Risk'],
    rows=[
        ['Traditional well', '275', '27.5%', 'High'],
        ['Improved well', '200', '20.0%', 'Moderate'],
        ['Borehole', '196', '19.6%', 'Low'],
        ['River / Spring', '167', '16.7%', 'Very high'],
        ['Other', '82', '8.2%', 'Undetermined'],
        ['SODECI (piped)', '80', '8.0%', 'Low'],
        ['Total', '1,000', '100%', '-'],
    ],
    caption='Table 3 - Water source details',
    col_widths=[5.5, 3.5, 3.0, 4.0]
)

add_heading(doc, '3.2 Water quality by department', 0, 'WASH_H2')
add_chart(doc, 'qualite_dept', width_inches=4.5,
          caption='Figure 3 - Perceived water quality by department')

add_chart(doc, 'traitement', width_inches=3.5,
          caption='Figure 4 - Household water treatment')

add_table_styled(doc,
    headers=['Department', 'Good', 'Acceptable', 'Poor', 'Total', '% poor'],
    rows=[
        ['Goh', '118 (20.7%)', '221 (38.8%)', '231 (40.5%)', '570', '40.5%'],
        ['Loh Djiboua', '90 (20.9%)', '171 (39.8%)', '169 (39.3%)', '430', '39.3%'],
        ['Total', '208 (20.8%)', '392 (39.2%)', '400 (40.0%)', '1,000', '40.0%'],
    ],
    caption='Table 4 - Water quality by department',
    col_widths=[3.5, 3.0, 3.0, 3.0, 2.0, 2.5]
)

add_heading(doc, '3.3 Water access and waterborne diseases', 0, 'WASH_H2')
add_chart(doc, 'acces', width_inches=4.5,
          caption='Figure 5 - Time to reach water source')

add_chart(doc, 'maladies', width_inches=4.5,
          caption='Figure 6 - Reported waterborne diseases by department')

add_table_styled(doc,
    headers=['Time bracket', 'Households', '%', 'WHO standard', 'Status'],
    rows=[
        ['Under 30 minutes', '386', '38.6%', 'Compliant', 'Acceptable'],
        ['30 to 60 minutes', '341', '34.1%', 'Borderline', 'Caution'],
        ['Over 60 minutes', '273', '27.3%', 'Non-compliant', 'Critical'],
    ],
    caption='Table 5 - Water access time and implications',
    col_widths=[4.5, 2.5, 2.5, 3.5, 3.0]
)

add_table_styled(doc,
    headers=['Disease', 'Goh', 'Loh Djiboua', 'Estimated total', 'Severity'],
    rows=[
        ['Diarrhea', '~115', '~87', '~202', 'Moderate'],
        ['Typhoid', '~72', '~55', '~127', 'High'],
        ['Hepatitis A', '~60', '~43', '~103', 'High'],
        ['Cholera', '~50', '~34', '~84', 'Very high'],
        ['Total reported', '~297', '~219', '~516 cases', '-'],
    ],
    caption='Table 6 - Reported waterborne diseases',
    col_widths=[4.0, 2.5, 3.0, 3.0, 3.5]
)

# ===== SECTION 4: SANITATION INFRASTRUCTURE =====
print("Section 4: Sanitation Infrastructure...")
add_page_break(doc)
add_heading(doc, 'Section 4 - Sanitation Infrastructure', 0, 'WASH_H1')

add_para(doc,
    "Sanitation represents a major challenge in both departments. Nearly one third "
    "of households (33.6%) have no access to latrines, and only 14.8% have improved latrines "
    "or modern toilets, far from SDG 6 targets.",
    'WASH_Body')

add_chart(doc, 'latrines', width_inches=5.0,
          caption='Figure 7 - Sanitation infrastructure types (N=1000)')

add_table_styled(doc,
    headers=['Infrastructure type', 'No. households', '%', 'SDG 6 compliance'],
    rows=[
        ['Simple latrine', '424', '42.4%', 'Partial'],
        ['No latrine (open defecation)', '336', '33.6%', 'Non-compliant'],
        ['Other type', '92', '9.2%', 'Undetermined'],
        ['Improved latrine', '90', '9.0%', 'Compliant'],
        ['Modern toilet', '58', '5.8%', 'Compliant'],
        ['Total', '1,000', '100%', '-'],
    ],
    caption='Table 7 - Latrine types and SDG 6 compliance',
    col_widths=[6.0, 3.0, 2.5, 4.5]
)

add_table_styled(doc,
    headers=['Indicator', 'Value', 'SDG 6 target', 'Gap'],
    rows=[
        ['Access to latrines (all types)', '66.4%', '100%', '-33.6 pts'],
        ['Improved latrines + modern toilets', '14.8%', '100%', '-85.2 pts'],
        ['Open defecation (OD)', '33.6%', '0%', '+33.6 pts'],
        ['Household water treatment', '29.1%', '100%', '-70.9 pts'],
    ],
    caption='Table 8 - Sanitation indicators vs SDG 6 targets',
    col_widths=[5.5, 3.0, 3.0, 4.5]
)

# ===== SECTION 5: WASTE MANAGEMENT =====
print("Section 5: Waste Management...")
add_page_break(doc)
add_heading(doc, 'Section 5 - Waste Management', 0, 'WASH_H1')

add_para(doc,
    "Waste management is a complex issue in both departments. "
    "52.3% of households are exposed to open dumps. Satisfaction remains mostly "
    "negative with 62% of households somewhat or completely unsatisfied with current management.",
    'WASH_Body')

add_chart(doc, 'gestion_modes', width_inches=5.0,
          caption='Figure 8 - Waste disposal methods')

add_chart(doc, 'dechets_sat', width_inches=5.0,
          caption='Figure 9 - Waste management satisfaction')

add_table_styled(doc,
    headers=['Disposal method', 'Households', '%', 'Environmental impact'],
    rows=[
        ['Collection / Recycling', '515', '51.5%', 'Positive'],
        ['Open dumping', '501', '50.1%', 'Very negative'],
        ['Burning', '499', '49.9%', 'Negative'],
        ['Reuse', '287', '28.7%', 'Positive'],
    ],
    caption='Table 9 - Waste disposal methods',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Satisfaction level', 'Households', '%', 'Trend'],
    rows=[
        ['Somewhat satisfied', '378', '37.8%', 'Negative'],
        ['Satisfied', '271', '27.1%', 'Positive'],
        ['Not satisfied', '242', '24.2%', 'Negative'],
        ['Very satisfied', '109', '10.9%', 'Positive'],
        ['Total', '1,000', '100%', '-'],
    ],
    caption='Table 10 - Waste management satisfaction',
    col_widths=[5.0, 3.0, 2.5, 5.5]
)

# ===== SECTION 6: ATMOSPHERIC POLLUTION =====
print("Section 6: Atmospheric Pollution...")
add_page_break(doc)
add_heading(doc, 'Section 6 - Atmospheric Pollution', 0, 'WASH_H1')

add_para(doc,
    "Atmospheric pollution is reported by nearly one household in two (49.4%). "
    "Air quality is judged poor or very poor by 51.6% of respondents. "
    "Noise nuisances also affect a majority of households (57.3%).",
    'WASH_Body')

add_chart(doc, 'pollution', width_inches=4.5,
          caption='Figure 10 - Reported atmospheric pollution by department')

add_chart(doc, 'qualite_air', width_inches=4.5,
          caption='Figure 11 - Perceived air quality (N=1000)')

add_table_styled(doc,
    headers=['Indicator', 'Households', '%', 'Risk level'],
    rows=[
        ['Atmospheric pollution reported', '494', '49.4%', 'Concerning'],
        ['Air quality: Poor', '364', '36.4%', 'Critical'],
        ['Air quality: Average', '361', '36.1%', 'Moderate'],
        ['Air quality: Very poor', '152', '15.2%', 'Very critical'],
        ['Air quality: Good', '123', '12.3%', 'Acceptable'],
        ['Noise nuisances', '573', '57.3%', 'Concerning'],
    ],
    caption='Table 11 - Pollution and air quality indicators',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Department', 'Pollution reported', 'No pollution', 'Rate'],
    rows=[
        ['Goh', '285', '285', '50.0%'],
        ['Loh Djiboua', '209', '221', '48.6%'],
        ['Total', '494', '506', '49.4%'],
    ],
    caption='Table 12 - Pollution by department',
    col_widths=[4.0, 4.0, 4.0, 4.0]
)

# ===== SECTION 7: NATURAL RESOURCES =====
print("Section 7: Natural Resources...")
add_page_break(doc)
add_heading(doc, 'Section 7 - Natural Resources', 0, 'WASH_H1')

add_para(doc,
    "The natural resources situation is alarming: 57.7% of households report "
    "resource degradation, and only 12.1% consider them abundant. "
    "This situation indicates growing awareness but practices that remain "
    "insufficient to reverse the trend.",
    'WASH_Body')

add_chart(doc, 'ressources', width_inches=5.0,
          caption='Figure 12 - Perceived natural resources availability (N=1000)')

add_table_styled(doc,
    headers=['Availability level', 'Households', '%', 'Status'],
    rows=[
        ['Scarce', '384', '38.4%', 'Critical'],
        ['Moderately available', '335', '33.5%', 'Precarious'],
        ['Near depletion', '160', '16.0%', 'Emergency'],
        ['Abundant', '121', '12.1%', 'Favorable'],
    ],
    caption='Table 13 - Natural resources availability',
    col_widths=[5.5, 3.0, 2.5, 5.0]
)

add_table_styled(doc,
    headers=['Resources indicator', 'Value', 'Interpretation'],
    rows=[
        ['Degradation reported', '57.7% (577)', 'Critical situation'],
        ['Environmental stewardship: Strong', '35.8% (358)', 'Awareness present'],
        ['Environmental stewardship: Moderate', '30.8% (308)', 'Needs reinforcement'],
        ['Environmental stewardship: Very strong', '18.3% (183)', 'Committed minority'],
        ['Environmental stewardship: Weak', '15.1% (151)', 'High risk'],
    ],
    caption='Table 14 - Natural resources and environmental stewardship indicators',
    col_widths=[5.5, 4.0, 6.5]
)

# ===== SECTION 8: COMPARATIVE ANALYSIS =====
print("Section 8: Comparative Analysis...")
add_page_break(doc)
add_heading(doc, 'Section 8 - Comparative Analysis', 0, 'WASH_H1')

add_para(doc,
    "The comparison of WASH indicators between the two departments reveals relatively similar "
    "profiles. The Goh department shows slightly more disease cases and a higher proportion of "
    "treated water, while Loh Djiboua shows better access within 30 minutes.",
    'WASH_Body')

add_chart(doc, 'comparaison', width_inches=5.5,
          caption='Figure 13 - Comparison of key WASH indicators by department')

add_table_styled(doc,
    headers=['Indicator', 'Overall', 'Goh (N=570)', 'Loh Djiboua (N=430)', 'Gap'],
    rows=[
        ['Good water quality', '20.8%', '20.7%', '20.9%', '+0.2 pts'],
        ['Poor water quality', '40.0%', '40.5%', '39.3%', '-1.2 pts'],
        ['Access <30 min', '38.6%', '37.5%', '40.0%', '+2.5 pts'],
        ['Access >60 min', '27.3%', '28.1%', '26.3%', '-1.8 pts'],
        ['Treated water', '29.1%', '30.0%', '28.1%', '-1.9 pts'],
        ['Waterborne diseases', '27.1%', '26.7%', '27.7%', '+1.0 pts'],
        ['No latrine (OD)', '33.6%', '33.3%', '33.3%', '0.0 pts'],
        ['Atmospheric pollution', '49.4%', '50.0%', '48.6%', '-1.4 pts'],
        ['Scarce resources', '38.4%', '~39.0%', '~37.5%', '-1.5 pts'],
        ['Resource degradation', '57.7%', '~58.0%', '~57.2%', '-0.8 pts'],
        ['Sufficient water', '58.0%', '~58.5%', '~57.2%', '-1.3 pts'],
        ['Exposed to waste', '52.3%', '~53.0%', '~51.4%', '-1.6 pts'],
    ],
    caption='Table 15 - Full comparative table of WASH indicators',
    col_widths=[4.0, 2.5, 3.0, 3.5, 3.0]
)

# ===== SECTION 9: PRIORITY RECOMMENDATIONS =====
print("Section 9: Priority Recommendations...")
add_page_break(doc)
add_heading(doc, 'Section 9 - Priority Recommendations', 0, 'WASH_H1')

add_para(doc,
    "Based on the findings of this survey, the following recommendations are formulated, "
    "ranked in order of priority according to urgency and potential impact on public health "
    "and the environment.",
    'WASH_Body')

recos = [
    ('Priority 1 - Water quality and treatment', [
        'Mass campaigns on household water treatment (chlorination, filtration, boiling)',
        'Construction of additional boreholes in high-risk areas',
        'Regular testing of water source quality',
        'Rehabilitation of existing traditional wells',
    ]),
    ('Priority 2 - Sanitation & Latrines', [
        'Open defecation free (ODF) elimination programme',
        'Subsidies for construction of improved latrines',
        'Awareness of waterborne diseases linked to lack of sanitation',
        'Training local masons in latrine construction',
    ]),
    ('Priority 3 - Waste management', [
        'Establish structured collection systems in underserved areas',
        'Promote composting and organic waste recycling',
        'Ban and penalise open-air burning',
        'Regulated and secured transit depots',
    ]),
    ('Priority 4 - Atmospheric pollution', [
        'Mapping and monitoring of pollution sources',
        'Regulation of industrial and artisanal emissions',
        'Promote improved cookstoves to reduce domestic smoke',
        'Buffer zones around urban agglomerations',
    ]),
    ('Priority 5 - Natural resources', [
        'Local sustainable natural resource management plans',
        'Reforestation and degraded land restoration',
        'Community awareness on environmental preservation',
        'Participatory resource monitoring systems',
    ]),
    ('Priority 6 - Monitoring & Governance', [
        'WASH indicator monitoring and evaluation systems',
        'Capacity building for local authorities',
        'Community involvement in infrastructure management',
        'Repeat the survey every 2 years to measure progress',
    ]),
]

for title, items in recos:
    add_heading(doc, title, 0, 'WASH_H2')
    for item in items:
        p = doc.add_paragraph(style='WASH_Bullet')
        run = p.add_run('   >> ' + item)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    doc.add_paragraph()

# ===== SECTION 10: CONCLUSION =====
print("Section 10: Conclusion...")
add_page_break(doc)
add_heading(doc, 'Section 10 - Conclusion', 0, 'WASH_H1')

add_para(doc,
    "This WASH survey conducted among 1,000 households in the Goh and "
    "Loh Djiboua departments of Cote d'Ivoire highlights considerable challenges in "
    "drinking water, sanitation and environmental management.",
    'WASH_Body')

add_para(doc,
    "The most concerning findings relate to water quality (40% poor quality), "
    "the low rate of household treatment (29.1%), and the prevalence of open defecation "
    "(33.6%). These combined factors largely explain the 27.1% of households reporting "
    "water-related diseases.",
    'WASH_Body')

add_para(doc,
    "Relative to the Sustainable Development Goals (SDG 6) - ensuring access to safe water "
    "and sanitation for all - both departments remain far from targets. Targeted investments "
    "and behaviour change programmes are essential to reduce access inequalities and protect "
    "population health.",
    'WASH_Body')

add_para(doc,
    "As the profiles of the two departments are very similar, interventions can be "
    "planned in a coordinated manner, while accounting for the local specificities of each "
    "sub-prefecture. A repeat survey in two years will allow progress to be measured "
    "and intervention strategies to be adapted.",
    'WASH_Body')

add_heading(doc, 'Summary of key indicators', 0, 'WASH_H2')
add_table_styled(doc,
    headers=['Dimension', 'Main indicator', 'Value', 'Priority'],
    rows=[
        ['Water', 'Poor quality', '40%', 'Critical'],
        ['Water', 'Untreated', '70.9%', 'Critical'],
        ['Access', 'Over 60 min', '27.3%', 'High'],
        ['Health', 'Waterborne diseases', '27.1%', 'High'],
        ['Sanitation', 'No latrine', '33.6%', 'Critical'],
        ['Waste', 'Somewhat/not satisfied', '62.0%', 'Moderate'],
        ['Pollution', 'Reported', '49.4%', 'Moderate'],
        ['Resources', 'Degradation', '57.7%', 'High'],
        ['Environment', 'Scarce resources', '38.4%', 'High'],
    ],
    caption='Table 16 - Summary of key indicators and priority levels',
    col_widths=[4.0, 5.5, 3.0, 3.5]
)

# ===== SAVE =====
out_path = r'C:\Users\ITrebi\Music\PortFolio Upwork\Excel Dashboard\WASH_Report_EN.docx'
doc.save(out_path)
print("Word report saved to: " + out_path)
